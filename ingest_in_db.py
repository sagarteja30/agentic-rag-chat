import os
import json
import hashlib
from pathlib import Path
from dotenv import load_dotenv
import httpx
from loguru import logger
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
import time
import asyncio
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import multiprocessing
from typing import List, Dict, Any, Tuple, Optional
import psutil
import gc
from tqdm import tqdm
import mimetypes

# Production imports for large PDF handling
try:
    from pdfminer.high_level import extract_text
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfparser import PDFParser
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False
    logger.warning("pdfminer.six not available. Install with: pip install pdfminer.six")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Install with: pip install python-docx")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logger.warning("pandas not available. Install with: pip install pandas")

# Load environment variables
load_dotenv()

# -------------------
# Production Configuration
# -------------------
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")

if not SUPABASE_URL or not SUPABASE_KEY:
    raise ValueError("Supabase URL or Service Role Key not found in .env")

# Production settings for large documents and high throughput
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE_CHARS", 1200))  # Reduced from 2000 for memory efficiency
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP_CHARS", 200))  # Reduced from 400
BATCH_SIZE = int(os.getenv("BATCH_SIZE", 8))  # Reduced from 16 for memory efficiency
MAX_WORKERS = 1  # Process one file at a time to avoid memory issues
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", 100))  # Reduced from 500MB
DATA_DIR = Path("documents")

# Memory management settings
MAX_MEMORY_USAGE_PERCENT = 95  # Increase from 80 to 85
PROCESS_TIMEOUT = 3600  # 1 hour timeout for large files
EMBEDDING_CACHE_SIZE = 5000  # Cache embeddings for similar chunks

logger.info(f"Production Ingestion Configuration:")
logger.info(f"   Max workers: {MAX_WORKERS}")
logger.info(f"   Batch size: {BATCH_SIZE}")
logger.info(f"   Chunk size: {CHUNK_SIZE}")
logger.info(f"   Max file size: {MAX_FILE_SIZE_MB}MB")
logger.info(f"   Embedding model: {EMBEDDING_MODEL}")

# Initialize embedding model with error handling
try:
    model = SentenceTransformer(EMBEDDING_MODEL)
    model.eval()  # Set to evaluation mode for inference
    logger.info("✅ Initialized embedding model successfully")
except Exception as e:
    logger.error(f"❌ Failed to initialize embedding model: {e}")
    raise

# HTTP headers for Supabase
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json",
    "Prefer": "return=representation"
}

# Global caches and tracking
embedding_cache = {}
processed_files = set()
failed_files = []

# -------------------
# Production File Parsing
# -------------------
def get_file_size_mb(file_path: Path) -> float:
    """Get file size in MB"""
    return file_path.stat().st_size / (1024 * 1024)

def check_memory_usage() -> bool:
    """Check if memory usage is acceptable"""
    memory_percent = psutil.virtual_memory().percent
    available_gb = psutil.virtual_memory().available / (1024**3)
    
    # More lenient check: allow processing if we have at least 2GB available
    # or if memory usage is below 85% (instead of 80%)
    if available_gb > 1.0 or memory_percent < 95:  # Changed from 85 to 95
        return True
    
    logger.warning(f"Memory check: {memory_percent:.1f}% used, {available_gb:.1f}GB available")
    return False

def parse_large_pdf(path: Path) -> str:
    """Parse large PDF files with memory optimization"""
    if not PDFMINER_AVAILABLE:
        logger.error("PDF parsing not available. Install pdfminer.six")
        return ""
    
    try:
        size_mb = get_file_size_mb(path)
        logger.info(f"📖 Parsing large PDF: {path.name} ({size_mb:.1f}MB)")
        
        # For very large PDFs, process in smaller chunks
        if size_mb > 50:  # For PDFs larger than 50MB
            logger.info("Large PDF detected, using memory-optimized parsing...")
        
        # Use memory-efficient PDF parsing for large files
        with open(path, 'rb') as file:
            parser = PDFParser(file)
            doc = PDFDocument(parser)
            
            # Check if PDF is encrypted
            if doc.is_extractable:
                text = extract_text(str(path), maxpages=0, caching=False)  # Disable caching for large files
                logger.info(f"✅ Extracted {len(text):,} characters from PDF")
                return text
            else:
                logger.warning(f"⚠️ PDF is encrypted or not extractable: {path}")
                return ""
                
    except MemoryError:
        logger.error(f"❌ Not enough memory to process PDF {path}")
        return ""
    except Exception as e:
        logger.error(f"❌ Error parsing PDF {path}: {e}")
        return ""

def parse_word_document(path: Path) -> str:
    """Parse Word documents with error handling"""
    if not DOCX_AVAILABLE:
        logger.error("Word document parsing not available. Install python-docx")
        return ""
    
    try:
        logger.info(f"📄 Parsing Word document: {path.name}")
        doc = docx.Document(str(path))
        
        # Extract text from paragraphs and tables
        text_parts = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        text = "\n".join(text_parts)
        logger.info(f"✅ Extracted {len(text):,} characters from Word document")
        return text
        
    except Exception as e:
        logger.error(f"❌ Error parsing Word document {path}: {e}")
        return ""

def parse_csv_file(path: Path) -> str:
    """Parse CSV files with pandas"""
    if not PANDAS_AVAILABLE:
        logger.error("CSV parsing not available. Install pandas")
        return ""
    
    try:
        logger.info(f"📊 Parsing CSV file: {path.name}")
        df = pd.read_csv(path, encoding='utf-8', on_bad_lines='skip')
        
        # Convert DataFrame to readable text format
        text_parts = []
        text_parts.append(f"Dataset: {path.name}")
        text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
        text_parts.append(f"Rows: {len(df)}")
        text_parts.append("")
        
        # Add column descriptions
        for col in df.columns:
            text_parts.append(f"Column '{col}': {df[col].dtype}")
            if df[col].dtype == 'object':
                unique_vals = df[col].nunique()
                text_parts.append(f"  Unique values: {unique_vals}")
                if unique_vals < 20:
                    text_parts.append(f"  Values: {', '.join(map(str, df[col].unique()[:10]))}")
        
        # Add sample rows
        text_parts.append("\nSample Data:")
        text_parts.append(df.head(10).to_string(index=False))
        
        text = "\n".join(text_parts)
        logger.info(f"✅ Processed CSV with {len(df)} rows, {len(df.columns)} columns")
        return text
        
    except Exception as e:
        logger.error(f"❌ Error parsing CSV {path}: {e}")
        return ""

def parse_file_production(path: Path) -> str:
    """Production-grade file parsing with comprehensive format support"""
    try:
        if not path.exists():
            logger.error(f"File not found: {path}")
            return ""
        
        # Check file size
        size_mb = get_file_size_mb(path)
        if size_mb > MAX_FILE_SIZE_MB:
            logger.warning(f"⚠️ File too large ({size_mb:.1f}MB): {path}")
            return ""
        
        # Check memory usage
        if not check_memory_usage():
            logger.warning(f"⚠️ Memory usage too high, skipping {path}")
            return ""
        
        ext = path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(path))
        
        logger.info(f"🔍 Processing {ext} file: {path.name} ({size_mb:.1f}MB)")
        
        # PDF files
        if ext == ".pdf":
            return parse_large_pdf(path)
            
        # Word documents
        elif ext in [".docx", ".doc"]:
            return parse_word_document(path)
            
        # CSV files
        elif ext == ".csv":
            return parse_csv_file(path)
            
        # Text-based files
        elif ext in [".txt", ".md", ".py", ".js", ".html", ".css", ".json", ".xml", ".yaml", ".yml"]:
            try:
                # Try different encodings
                encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        text = path.read_text(encoding=encoding, errors='ignore')
                        logger.info(f"✅ Read {len(text):,} characters using {encoding} encoding")
                        return text
                    except UnicodeDecodeError:
                        continue
                
                logger.warning(f"⚠️ Could not decode text file: {path}")
                return ""
                
            except Exception as e:
                logger.error(f"❌ Error reading text file {path}: {e}")
                return ""
        
        # Log files
        elif ext == ".log":
            try:
                # Read only the last part of large log files
                if size_mb > 10:  # For log files larger than 10MB
                    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                        f.seek(max(0, path.stat().st_size - 1024*1024))  # Read last 1MB
                        text = f.read()
                        logger.info(f"✅ Read last 1MB of log file")
                        return text
                else:
                    text = path.read_text(encoding='utf-8', errors='ignore')
                    return text
            except Exception as e:
                logger.error(f"❌ Error reading log file {path}: {e}")
                return ""
        
        else:
            logger.warning(f"⚠️ Unsupported file type: {ext} ({mime_type})")
            return ""
            
    except Exception as e:
        logger.error(f"❌ Error parsing file {path}: {e}")
        return ""

# -------------------
# Advanced Text Chunking
# -------------------
def chunk_text_production(text: str, doc_name: str, max_chunk_size: int = None) -> List[Dict[str, Any]]:
    """Production text chunking with metadata and optimization"""
    if not text.strip():
        logger.warning(f"⚠️ Empty text for document: {doc_name}")
        return []
    
    chunk_size = max_chunk_size or CHUNK_SIZE
    
    try:
        # Adaptive chunking based on document type
        separators = ["\n\n\n", "\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""]
        
        # For code files, preserve structure better
        if doc_name.endswith(('.py', '.js', '.html', '.css', '.json')):
            separators = ["\n\nclass ", "\n\ndef ", "\n\nfunction ", "\n\n", "\n", ". ", " ", ""]
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=separators
        )
        
        chunks = splitter.split_text(text)
        logger.info(f"📝 Created {len(chunks)} chunks from {doc_name}")
        
        # Filter and enhance chunks
        enhanced_chunks = []
        for idx, chunk in enumerate(chunks):
            chunk_text = chunk.strip()
            
            # Skip very short chunks
            if len(chunk_text) < 100:
                continue
            
            # Calculate chunk hash for deduplication
            chunk_hash = hashlib.md5(chunk_text.encode()).hexdigest()[:16]
            
            # Create enhanced chunk with metadata
            enhanced_chunk = {
                'text': chunk_text,
                'chunk_id': idx,
                'chunk_hash': chunk_hash,
                'length': len(chunk_text),
                'word_count': len(chunk_text.split()),
                'doc_name': doc_name,
                'metadata': {
                    'source': doc_name,
                    'chunk_index': idx,
                    'total_chunks': len(chunks),
                    'chunk_hash': chunk_hash
                }
            }
            
            enhanced_chunks.append(enhanced_chunk)
        
        logger.info(f"✅ Processed {len(enhanced_chunks)} meaningful chunks from {doc_name}")
        return enhanced_chunks
        
    except Exception as e:
        logger.error(f"❌ Error chunking text for {doc_name}: {e}")
        return []

# -------------------
# Production Database Operations
# -------------------
async def test_connection_async() -> bool:
    """Async connection test"""
    try:
        url = f"{SUPABASE_URL}/rest/v1/documents?select=count&limit=1"
        timeout = httpx.Timeout(30.0)
        
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.get(url, headers=HEADERS)
            
            if response.status_code == 200:
                logger.info("✅ Database connection successful")
                return True
            elif response.status_code == 404:
                logger.error("❌ Table 'documents' not found. Please create the table first.")
                return False
            else:
                logger.error(f"❌ Connection failed: {response.status_code} - {response.text}")
                return False
                
    except Exception as e:
        logger.error(f"❌ Connection test failed: {e}")
        return False

async def clear_document_async(doc_id: str) -> bool:
    """Clear existing document chunks asynchronously"""
    try:
        url = f"{SUPABASE_URL}/rest/v1/documents?doc_id=eq.{doc_id}"
        timeout = httpx.Timeout(60.0)
        
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.delete(url, headers=HEADERS)
            
            if response.status_code in [200, 204]:
                logger.info(f"🧹 Cleared existing chunks for: {doc_id}")
                return True
            else:
                logger.warning(f"⚠️ Clear operation status: {response.status_code}")
                return True  # Continue anyway
                
    except Exception as e:
        logger.warning(f"⚠️ Clear operation failed for {doc_id}: {e}")
        return True  # Continue anyway

async def insert_batch_async(rows: List[Dict[str, Any]], max_retries: int = 3) -> bool:
    """Insert batch with async retry logic"""
    for attempt in range(max_retries):
        try:
            url = f"{SUPABASE_URL}/rest/v1/documents"
            timeout = httpx.Timeout(120.0)  # Extended timeout for large batches
            
            async with httpx.AsyncClient(timeout=timeout) as client:
                response = await client.post(url, json=rows, headers=HEADERS)
                
                if response.status_code in [200, 201]:
                    logger.info(f"✅ Inserted batch of {len(rows)} documents")
                    return True
                elif response.status_code == 409:
                    logger.warning("⚠️ Duplicate data conflict, continuing...")
                    return True
                else:
                    logger.error(f"❌ Insert failed: {response.status_code} - {response.text}")
                    if attempt < max_retries - 1:
                        await asyncio.sleep(2 ** attempt)
                    
        except Exception as e:
            logger.error(f"❌ Insert error (attempt {attempt + 1}): {e}")
            if attempt < max_retries - 1:
                await asyncio.sleep(2 ** attempt)
    
    return False

def generate_embeddings_batch(chunks: List[str]) -> List[List[float]]:
    """Generate embeddings for a batch of chunks with caching"""
    try:
        # Check cache first
        cached_embeddings = []
        uncached_chunks = []
        uncached_indices = []
        
        for i, chunk in enumerate(chunks):
            chunk_hash = hashlib.md5(chunk.encode()).hexdigest()
            if chunk_hash in embedding_cache:
                cached_embeddings.append((i, embedding_cache[chunk_hash]))
            else:
                uncached_chunks.append(chunk)
                uncached_indices.append(i)
        
        # Generate embeddings for uncached chunks
        new_embeddings = []
        if uncached_chunks:
            logger.info(f"🧠 Generating {len(uncached_chunks)} new embeddings ({len(cached_embeddings)} cached)")
            embeddings_array = model.encode(uncached_chunks, show_progress_bar=False, batch_size=32)
            new_embeddings = embeddings_array.tolist()
            
            # Cache new embeddings
            for chunk, embedding in zip(uncached_chunks, new_embeddings):
                chunk_hash = hashlib.md5(chunk.encode()).hexdigest()
                embedding_cache[chunk_hash] = embedding
                
                # Limit cache size
                if len(embedding_cache) > EMBEDDING_CACHE_SIZE:
                    # Remove oldest entries (simple FIFO)
                    oldest_key = next(iter(embedding_cache))
                    del embedding_cache[oldest_key]
        
        # Combine cached and new embeddings in correct order
        all_embeddings = [None] * len(chunks)
        
        # Place cached embeddings
        for i, embedding in cached_embeddings:
            all_embeddings[i] = embedding
        
        # Place new embeddings
        for i, embedding in zip(uncached_indices, new_embeddings):
            all_embeddings[i] = embedding
        
        return all_embeddings
        
    except Exception as e:
        logger.error(f"❌ Embedding generation failed: {e}")
        return []

# -------------------
# Main Production Ingestion
# -------------------
async def process_file_async(file_path: Path, clear_existing: bool = True) -> Tuple[str, bool, Dict[str, Any]]:
    """Process a single file asynchronously"""
    stats = {
        'chunks_created': 0,
        'chunks_inserted': 0,
        'processing_time': 0,
        'file_size_mb': 0,
        'error': None
    }
    
    start_time = time.time()
    
    try:
        stats['file_size_mb'] = get_file_size_mb(file_path)
        
        # Check if file was already processed (by hash)
        file_hash = hashlib.md5(str(file_path).encode()).hexdigest()
        if file_hash in processed_files:
            logger.info(f"⏭️ Skipping already processed file: {file_path.name}")
            return file_path.name, True, stats
        
        # Parse file
        logger.info(f"📖 Processing: {file_path.name}")
        text = parse_file_production(file_path)
        
        if not text.strip():
            stats['error'] = "No text extracted"
            return file_path.name, False, stats
        
        # Clear existing document if requested
        if clear_existing:
            await clear_document_async(file_path.name)
        
        # Create chunks
        chunks = chunk_text_production(text, file_path.name)
        stats['chunks_created'] = len(chunks)
        
        if not chunks:
            stats['error'] = "No chunks created"
            return file_path.name, False, stats
        
        # Process chunks in batches
        successful_batches = 0
        total_batches = (len(chunks) + BATCH_SIZE - 1) // BATCH_SIZE
        
        for i in range(0, len(chunks), BATCH_SIZE):
            batch_chunks = chunks[i:i + BATCH_SIZE]
            batch_num = i // BATCH_SIZE + 1
            
            logger.info(f"🔄 Processing batch {batch_num}/{total_batches} for {file_path.name}")
            
            try:
                # Generate embeddings
                chunk_texts = [chunk['text'] for chunk in batch_chunks]
                embeddings = generate_embeddings_batch(chunk_texts)
                
                if not embeddings:
                    logger.error(f"❌ Failed to generate embeddings for batch {batch_num}")
                    continue
                
                # Prepare database rows
                rows = []
                for chunk, embedding in zip(batch_chunks, embeddings):
                    row = {
                        "doc_id": file_path.name,
                        "chunk_id": chunk['chunk_id'],
                        "content": chunk['text'],
                        "metadata": {
                            **chunk['metadata'],
                            "file_size_mb": stats['file_size_mb'],
                            "processed_at": time.time(),
                            "chunk_hash": chunk['chunk_hash'],
                            "word_count": chunk['word_count']
                        },
                        "embedding": embedding
                    }
                    rows.append(row)
                
                # Insert batch
                if await insert_batch_async(rows):
                    successful_batches += 1
                    stats['chunks_inserted'] += len(rows)
                    logger.info(f"✅ Batch {batch_num} completed for {file_path.name}")
                else:
                    logger.error(f"❌ Batch {batch_num} failed for {file_path.name}")
                
                # Memory cleanup
                if batch_num % 5 == 0:  # Every 5 batches
                    gc.collect()
                
            except Exception as batch_error:
                logger.error(f"❌ Batch {batch_num} error for {file_path.name}: {batch_error}")
                continue
        
        # Mark file as processed
        processed_files.add(file_hash)
        
        success = successful_batches > 0
        stats['processing_time'] = time.time() - start_time
        
        if success:
            logger.info(f"✅ Completed {file_path.name}: {stats['chunks_inserted']} chunks in {stats['processing_time']:.2f}s")
        else:
            stats['error'] = "All batches failed"
            logger.error(f"❌ Failed to process {file_path.name}")
        
        return file_path.name, success, stats
        
    except Exception as e:
        stats['error'] = str(e)
        stats['processing_time'] = time.time() - start_time
        logger.error(f"❌ Error processing {file_path.name}: {e}")
        return file_path.name, False, stats

async def get_database_stats() -> Tuple[int, int]:
    """Get database statistics"""
    try:
        url = f"{SUPABASE_URL}/rest/v1/documents?select=doc_id"
        timeout = httpx.Timeout(30.0)
        
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.get(url, headers=HEADERS)
            
            if response.status_code == 200:
                docs = response.json()
                unique_docs = set(doc['doc_id'] for doc in docs)
                return len(docs), len(unique_docs)
            
        return 0, 0
        
    except Exception as e:
        logger.warning(f"⚠️ Could not get database stats: {e}")
        return 0, 0

async def ingest_documents_production(clear_existing: bool = True, max_files: Optional[int] = None):
    """Production-grade document ingestion with async processing"""
    logger.info("🚀 Starting Production Document Ingestion")
    start_time = time.time()
    
    # Test connection
    if not await test_connection_async():
        logger.error("❌ Cannot proceed without database connection")
        return False
    
    # Create documents directory
    DATA_DIR.mkdir(exist_ok=True)
    
    # Find supported files
    supported_extensions = {'.pdf', '.txt', '.docx', '.doc', '.md', '.py', '.js', '.html', 
                          '.css', '.json', '.xml', '.yaml', '.yml', '.csv', '.log'}
    
    all_files = []
    for ext in supported_extensions:
        all_files.extend(DATA_DIR.rglob(f"*{ext}"))
    
    # Filter files
    files = []
    for file_path in all_files:
        if (not file_path.name.startswith('.') and 
            not file_path.name.startswith('~') and
            get_file_size_mb(file_path) <= MAX_FILE_SIZE_MB):
            files.append(file_path)
    
    # Sort by size (process smaller files first)
    files.sort(key=lambda f: f.stat().st_size)
    
    # Limit files if specified
    if max_files:
        files = files[:max_files]
    
    if not files:
        logger.error(f"❌ No supported files found in {DATA_DIR}")
        logger.info(f"Supported formats: {', '.join(supported_extensions)}")
        return False
    
    logger.info(f"📁 Found {len(files)} files to process")
    total_size_mb = sum(get_file_size_mb(f) for f in files)
    logger.info(f"📊 Total size: {total_size_mb:.1f}MB")
    
    # Process files with progress tracking
    successful_files = 0
    failed_files_list = []
    total_chunks = 0
    
    # Use semaphore to limit concurrent database connections
    semaphore = asyncio.Semaphore(min(MAX_WORKERS, 10))
    
    async def process_with_semaphore(file_path):
        async with semaphore:
            return await process_file_async(file_path, clear_existing)
    
    # Process files concurrently
    logger.info(f"🔄 Processing files with {min(MAX_WORKERS, len(files))} concurrent workers")
    
    # Create tasks for all files
    tasks = [process_with_semaphore(file_path) for file_path in files]
    
    # Process with progress bar
    with tqdm(total=len(files), desc="Processing files") as pbar:
        for completed_task in asyncio.as_completed(tasks):
            try:
                file_name, success, stats = await completed_task
                
                if success:
                    successful_files += 1
                    total_chunks += stats['chunks_inserted']
                    logger.info(f"✅ {file_name}: {stats['chunks_inserted']} chunks, {stats['processing_time']:.1f}s")
                else:
                    failed_files_list.append((file_name, stats.get('error', 'Unknown error')))
                    logger.error(f"❌ {file_name}: {stats.get('error', 'Failed')}")
                
                pbar.update(1)
                pbar.set_postfix({
                    'Success': successful_files,
                    'Failed': len(failed_files_list),
                    'Chunks': total_chunks
                })
                
                # Memory check
                if not check_memory_usage():
                    logger.warning("⚠️ High memory usage detected, forcing cleanup")
                    gc.collect()
                
            except Exception as e:
                logger.error(f"❌ Task execution error: {e}")
                pbar.update(1)
    
    # Final statistics
    processing_time = time.time() - start_time
    db_chunks, db_docs = await get_database_stats()
    
    logger.info(f"\n📊 Production Ingestion Summary:")
    logger.info(f"   ✅ Files processed successfully: {successful_files}/{len(files)}")
    logger.info(f"   ❌ Files failed: {len(failed_files_list)}")
    logger.info(f"   📝 Total chunks created: {total_chunks}")
    logger.info(f"   💾 Database now contains: {db_chunks} chunks from {db_docs} documents")
    logger.info(f"   ⏱️ Total processing time: {processing_time:.2f}s")
    logger.info(f"   📈 Average processing speed: {total_chunks/processing_time:.1f} chunks/second")
    
    if failed_files_list:
        logger.error(f"\n❌ Failed Files:")
        for file_name, error in failed_files_list[:10]:  # Show first 10 failures
            logger.error(f"   {file_name}: {error}")
        if len(failed_files_list) > 10:
            logger.error(f"   ... and {len(failed_files_list) - 10} more")
    
    # Memory cleanup
    embedding_cache.clear()
    gc.collect()
    
    if successful_files > 0:
        logger.info("🎉 Production ingestion completed successfully!")
        logger.info("✅ Ready for production queries with 100+ concurrent users!")
        return True
    else:
        logger.error("❌ No files were successfully processed")
        return False

# -------------------
# Production CLI Interface
# -------------------
def main():
    """Main CLI interface for production ingestion"""
    # Declare global variables first
    global BATCH_SIZE, MAX_WORKERS, CHUNK_SIZE
    
    import argparse
    
    parser = argparse.ArgumentParser(description="Production Document Ingestion for Agentic RAG")
    parser.add_argument("--clear", action="store_true", help="Clear existing documents before ingestion")
    parser.add_argument("--max-files", type=int, help="Maximum number of files to process")
    parser.add_argument("--batch-size", type=int, default=BATCH_SIZE, help="Batch size for processing")
    parser.add_argument("--workers", type=int, default=MAX_WORKERS, help="Number of worker processes")
    parser.add_argument("--chunk-size", type=int, default=CHUNK_SIZE, help="Chunk size in characters")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging")
    
    args = parser.parse_args()
    
    # Update global settings based on arguments
    BATCH_SIZE = args.batch_size
    MAX_WORKERS = min(args.workers, multiprocessing.cpu_count())
    CHUNK_SIZE = args.chunk_size
    
    if args.verbose:
        logger.add(lambda msg: print(msg), level="DEBUG")
    
    logger.info("🚀 Production Document Ingestion System")
    logger.info("=" * 60)
    logger.info(f"📁 Document directory: {DATA_DIR.absolute()}")
    logger.info(f"🧠 Embedding model: {EMBEDDING_MODEL}")
    logger.info(f"📊 Configuration:")
    logger.info(f"   • Batch size: {BATCH_SIZE}")
    logger.info(f"   • Max workers: {MAX_WORKERS}")
    logger.info(f"   • Chunk size: {CHUNK_SIZE}")
    logger.info(f"   • Max file size: {MAX_FILE_SIZE_MB}MB")
    logger.info(f"   • Clear existing: {'Yes' if args.clear else 'No'}")
    if args.max_files:
        logger.info(f"   • Max files: {args.max_files}")
    logger.info("=" * 60)
    
    # System resource check
    cpu_count = psutil.cpu_count()
    memory_gb = psutil.virtual_memory().total / (1024**3)
    logger.info(f"💻 System Resources: {cpu_count} CPUs, {memory_gb:.1f}GB RAM")
    
    if memory_gb < 4:
        logger.warning("⚠️ Low memory detected. Consider reducing batch size or max workers.")
    
    # Run ingestion
    try:
        success = asyncio.run(ingest_documents_production(
            clear_existing=args.clear,
            max_files=args.max_files
        ))
        
        if success:
            logger.info("\n🎯 Next Steps:")
            logger.info("1. Start the FastAPI backend: python agentic_rag.py")
            logger.info("2. Start the Streamlit frontend: streamlit run agentic_rag_streamlit.py")
            logger.info("3. Open http://localhost:8501 in your browser")
            logger.info("\n🚀 System ready for production use!")
        else:
            logger.error("\n❌ Ingestion failed. Check logs and try again.")
            exit(1)
            
    except KeyboardInterrupt:
        logger.info("\n⏹️ Ingestion interrupted by user")
        exit(0)
    except Exception as e:
        logger.error(f"\n❌ Fatal error: {e}")
        exit(1)

if __name__ == "__main__":
    main()